import io
import os
import json
from datetime import datetime
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from openai import OpenAI
from pypdf import PdfReader
import streamlit as st

# Configuração da página
st.set_page_config(
    page_title="AuditGuard SST — Plataforma de Inteligência",
    page_icon="🛡️",
    layout="wide",
)

# Estilização de Fundo Corporativo e Visual do App via CSS
st.markdown(
    """
    <style>
    .stApp {
        background: linear-gradient(rgba(248, 249, 250, 0.93), rgba(248, 249, 250, 0.93)),
                    url('https://images.unsplash.com/photo-1486406146926-c627a92ad1ab?q=80&w=2070');
        background-size: cover;
        background-attachment: fixed;
    }
    h1 {
        color: #003366 !important;
        font-weight: 700 !important;
    }
    .stButton>button {
        border-radius: 8px;
        background-color: #003366;
        color: white;
        font-weight: bold;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Persistência da sessão
if "docx_bytes_1" not in st.session_state:
    st.session_state.docx_bytes_1 = None
if "texto_resposta_1" not in st.session_state:
    st.session_state.texto_resposta_1 = None

if "docx_bytes_2" not in st.session_state:
    st.session_state.docx_bytes_2 = None
if "texto_resposta_2" not in st.session_state:
    st.session_state.texto_resposta_2 = None

if "res_2_text" not in st.session_state:
    st.session_state.res_2_text = ""

# ==========================================================================
# GERENCIAMENTO DE USUÁRIOS E COTAS (PROTEÇÃO DE TOKENS)
# ==========================================================================
_DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
_USERS_FILE = os.path.join(_DATA_DIR, "usuarios.json")
LIMITE_MENSAL_PADRAO = 200

def _carregar_usuarios() -> dict:
    os.makedirs(_DATA_DIR, exist_ok=True)
    if not os.path.exists(_USERS_FILE):
        return {}
    with open(_USERS_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
        mes_atual = datetime.now().month
        mudou = False
        for user in data:
            if data[user].get("mes_referencia") != mes_atual:
                data[user]["auditorias_mes_atual"] = 0
                data[user]["mes_referencia"] = mes_atual
                mudou = True
        if mudou:
            with open(_USERS_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        return data

def _salvar_usuarios(data: dict) -> None:
    os.makedirs(_DATA_DIR, exist_ok=True)
    with open(_USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def verificar_limite_auditorias(usuario: str) -> tuple[bool, int, int]:
    usuarios = _carregar_usuarios()
    user_data = usuarios.get(usuario, {})
    limite = user_data.get("limite_mensal", LIMITE_MENSAL_PADRAO)
    uso_atual = user_data.get("auditorias_mes_atual", 0)
    return (uso_atual < limite), uso_atual, limite

def incrementar_uso_auditoria(usuario: str):
    usuarios = _carregar_usuarios()
    if usuario in usuarios:
        usuarios[usuario]["auditorias_mes_atual"] = usuarios[usuario].get("auditorias_mes_atual", 0) + 1
        _salvar_usuarios(usuarios)

def extrair_texto_pdf(pdf_file):
    try:
        reader = PdfReader(pdf_file)
        texto = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                texto += t + "\n"
        return texto
    except Exception as e:
        st.error(f"Erro ao ler PDF: {e}")
        return ""

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def converter_markdown_para_word_paisagem(texto_markdown):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(9.5)

    linhas = texto_markdown.split("\n")
    tabela_atual = None
    row_count = 0

    for linha in linhas:
        linha_str = linha.strip()
        if not linha_str:
            continue

        if linha_str.startswith("# "):
            titulo_limpo = linha_str.replace("# ", "").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(titulo_limpo)
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 51, 102)

        elif linha_str.startswith("##"):
            sub_limpo = linha_str.lstrip("#").strip()
            p = doc.add_paragraph()
            run = p.add_run(sub_limpo)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 102, 153)

        elif linha_str.startswith("|") and linha_str.endswith("|"):
            colunas = [c.strip() for c in linha_str.split("|")[1:-1]]
            if all(set(c) <= set("-: ") for c in colunas):
                continue

            if tabela_atual is None:
                row_count = 0
                tabela_atual = doc.add_table(rows=1, cols=len(colunas))
                tabela_atual.style = "Table Grid"
                hdr_cells = tabela_atual.rows[0].cells
                for i, col_texto in enumerate(colunas):
                    hdr_cells[i].text = col_texto
                    set_cell_background(hdr_cells[i], "003366")
                    for p in hdr_cells[i].paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                row_count += 1
                row_cells = tabela_atual.add_row().cells
                cor_fundo = "F2F4F8" if row_count % 2 == 0 else "FFFFFF"
                for i, col_texto in enumerate(colunas):
                    if i < len(row_cells):
                        row_cells[i].text = col_texto
                        set_cell_background(row_cells[i], cor_fundo)
                        for p in row_cells[i].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(8.5)
        else:
            tabela_atual = None
            p = doc.add_paragraph()
            partes = linha_str.split("**")
            for idx, parte in enumerate(partes):
                run = p.add_run(parte)
                if idx % 2 != 0:
                    run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Sidebar
LOGO_URL = "https://raw.githubusercontent.com/junioramaral0112/AnalisePGR/main/logo.png"
st.sidebar.image(LOGO_URL, width=220)
st.sidebar.title("⚙️ Painel do Sistema")
st.sidebar.info("**AuditGuard SST**\n\nEcossistema integrado de Inteligência Artificial para SST.")

usuario_logado = st.session_state.get("usuario", "admin")
permitido_sidebar, uso_sidebar, limite_sidebar = verificar_limite_auditorias(usuario_logado)
st.sidebar.markdown("---")
st.sidebar.markdown("### 📊 Cota Mensal (IA)")
st.sidebar.progress(min(uso_sidebar / limite_sidebar, 1.0))
st.sidebar.write(f"Utilizadas: **{uso_sidebar} / {limite_sidebar}**")

# Cabeçalho Principal
col_logo, col_titulo = st.columns([1, 5])
with col_logo:
    st.image(LOGO_URL, width=110)
with col_titulo:
    st.title("AuditGuard SST — Central de Inteligência")
    st.markdown("**Selecione o módulo de auditoria desejado nas abas abaixo:**")

st.markdown("---")

# ABAS DO SISTEMA
aba_completa, aba_pcmso = st.tabs(["🚀 Auditoria Completa (PGR + PCMSO + Apêndices)", "🩺 Auditoria Focada (PCMSO & ASO)"])

# ==========================================
# ABA 1: AUDITORIA COMPLETA ORIGINAL
# ==========================================
with aba_completa:
    st.markdown("### Módulo de Contestação Geral e Apêndices")
    st.markdown("Carregue os arquivos do **PGR** e **PCMSO** e informe as ressalvas gerais da auditoria.")

    col1, col2 = st.columns(2)
    with col1:
        pgr_file_1 = st.file_uploader("📄 Upload do PGR (PDF):", type=["pdf"], key="pgr_1")
    with col2:
        pcmso_file_1 = st.file_uploader("🩺 Upload do PCMSO (PDF):", type=["pdf"], key="pcmso_1")

    ressalvas_text_1 = st.text_area(
        "📝 Cole o texto das Ressalvas / Glosas Gerais:",
        height=140,
        placeholder="Exemplo: O documento foi Aceito com Ressalva pois os riscos psicossociais não constam no Inventário...",
        key="res_1"
    )

    btn_analisar_1 = st.button("🚀 Executar Auditoria Completa", key="btn_1")

    if btn_analisar_1:
        permitido, uso_atual, limite_max = verificar_limite_auditorias(usuario_logado)
        if not permitido:
            st.error(f"🚫 **Limite mensal atingido:** {limite_max} auditorias permitidas para este mês.")
            st.stop()

        api_key = str(st.secrets.get("DEEPSEEK_API_KEY", os.environ.get("DEEPSEEK_API_KEY", ""))).strip()

        if not api_key:
            st.error("⚠️ Chave DeepSeek não configurada!")
        elif not pgr_file_1 or not pcmso_file_1:
            st.error("⚠️ Envie ambos os arquivos (PGR e PCMSO)!")
        elif not ressalvas_text_1.strip():
            st.error("⚠️ Digite as ressalvas!")
        else:
            with st.spinner("⏳ Processando auditoria completa e gerando Word..."):
                try:
                    texto_pgr = extrair_texto_pdf(pgr_file_1)
                    texto_pcmso = extrair_texto_pdf(pcmso_file_1)
                    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

                    prompt_1 = f"""
                    Você é um Engenheiro de Segurança do Trabalho e Médico do Trabalho Sênior.
                    Emita um Parecer Técnico de Contestação de SST e a revisão integral dos Apêndices A e B.

                    --- PGR ---
                    {texto_pgr[:35000]}

                    --- PCMSO ---
                    {texto_pcmso[:35000]}

                    --- RESSALVAS ---
                    {ressalvas_text_1}

                    ESTRUTURA:
                    # PARECER TÉCNICO DE CONTESTAÇÃO DE SST — AUDITGUARD
                    1. **Objeto e Finalidade:** Descrição completa.
                    2. **Análise Crítica das Ressalvas:** Procedente ou Improcedente.
                    3. **Fundamentação Legal:** NR-01, NR-07, NR-17.

                    ## APÊNDICE A – INVENTÁRIO DE RISCOS OCUPACIONAIS CONSOLIDADO
                    | Tipo de Risco | Exposição | Perigo / Fonte Geradora | Resultado da Avaliação | Reconhecido? | Risco Ocupacional / Dano | Avaliação Inicial (P x S) | Medidas de Controle / EPIs / EPCs | Avaliação Residual (P x S) | eSocial / Plano de Ação |

                    ## APÊNDICE B – PLANO DE AÇÃO DE SAÚDE MENTAL E RISCOS PSICOSSOCIAIS
                    | Ação Preventiva / Corretiva | Fator de Risco | Responsável | Prazo | Indicador de Acompanhamento |
                    """

                    response = client.chat.completions.create(
                        model="deepseek-chat",
                        messages=[{"role": "system", "content": "Motor de IA AuditGuard."}, {"role": "user", "content": prompt_1}],
                        stream=False,
                    )

                    resp_text = response.choices[0].message.content
                    bytes_word = converter_markdown_para_word_paisagem(resp_text)
                    incrementar_uso_auditoria(usuario_logado)

                    st.session_state.texto_resposta_1 = resp_text
                    st.session_state.docx_bytes_1 = bytes_word
                    st.success("✅ Auditoria Completa gerada com sucesso!")
                except Exception as e:
                    st.error(f"❌ Erro: {str(e)}")

    if st.session_state.docx_bytes_1 and st.session_state.texto_resposta_1:
        st.markdown("---")
        st.download_button("📥 Baixar Parecer Completo (.docx)", data=st.session_state.docx_bytes_1, file_name="AuditGuard_Completo.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_1")
        st.write(st.session_state.texto_resposta_1)

# ==========================================
# ABA 2: AUDITORIA FOCADA (PCMSO & ASO)
# ==========================================
with aba_pcmso:
    st.markdown("### Módulo de Auditoria Médica e Cruzamento com ASO")
    st.markdown("Carregue o **PCMSO** e o **ASO**, para verificar se os exames realizados atendem rigorosamente às diretrizes da NR-07.")

    col1, col2 = st.columns(2)
    with col1:
        pcmso_file_2 = st.file_uploader("🩺 Upload do PCMSO (PDF):", type=["pdf"], key="pcmso_2")
    with col2:
        aso_file_2 = st.file_uploader("📋 Upload do ASO (PDF):", type=["pdf"], key="aso_2")

    if st.button("💡 Usar Exemplo Padrão de Auditoria ASO x PCMSO"):
        st.session_state.res_2_text = "Verificar se o ASO está em conformidade com o PCMSO, se os exames realizados são exatamente os requeridos pelo PCMSO para a função e se há exames faltantes."

    ressalvas_text_2 = st.text_area(
        "📝 Cole ou digite as Ressalvas / Glosas Médicas:",
        value=st.session_state.res_2_text,
        height=140,
        placeholder="Exemplo: Verificar se o ASO está em conformidade com o PCMSO, se os exames são os requeridos...",
        key="res_2"
    )

    btn_analisar_2 = st.button("🚀 Executar Auditoria Médica (PCMSO/ASO)", key="btn_2")

    if btn_analisar_2:
        permitido, uso_atual, limite_max = verificar_limite_auditorias(usuario_logado)
        if not permitido:
            st.error(f"🚫 **Limite mensal atingido:** {limite_max} auditorias permitidas.")
            st.stop()

        api_key = str(st.secrets.get("DEEPSEEK_API_KEY", os.environ.get("DEEPSEEK_API_KEY", ""))).strip()

        if not api_key:
            st.error("⚠️ Chave DeepSeek não configurada!")
        elif not pcmso_file_2 or not aso_file_2:
            st.error("⚠️ É obrigatório enviar ambos os arquivos (PCMSO e ASO) para realizar a auditoria cruzada!")
        elif not ressalvas_text_2.strip():
            st.error("⚠️ Digite as ressalvas médicas!")
        else:
            with st.spinner("⏳ Analisando convergência entre PCMSO e ASO e gerando parecer médico em Word..."):
                try:
                    texto_pcmso = extrair_texto_pdf(pcmso_file_2)
                    texto_aso = extrair_texto_pdf(aso_file_2)
                    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

                    prompt_2 = f"""
                    Você é um Médico do Trabalho Sênior, especialista em NR-07 e eSocial (S-2220).
                    
                    ATENÇÃO CRÍTICA DE LEITURA: Os documentos de ASO muitas vezes possuem campos preenchidos à mão ou em formato de formulário digitalizado. Faça uma varredura extremamente cuidadosa no texto extraído do ASO para identificar a seção "Exames Realizados e Data de Realização". 
                    NÃO assuma cegamente que exames estão ausentes se houver menção a eles no texto (como Acuidade Visual, Psicossocial, Hemograma, Glicemia, Reticulócitos, Exame Clínico, etc.). Só aponte ausência caso o exame realmente não conste no documento.

                    Emita um Parecer Técnico Médico focado estritamente em apontar de forma cirúrgica e detalhada se o ASO enviado está ou não em conformidade com o PCMSO.

                    ATENÇÃO CRÍTICA: A seção de "Conclusão e Providências" deve trazer o **Veredito Geral** em absoluto destaque (utilizando formatação clara e enfática em negrito/caixa alta), indicando sem margem de dúvida se o ASO está ou não aprovado em conformidade com o PCMSO fornecido com base nos dados reais encontrados.

                    --- PCMSO ---
                    {texto_pcmso[:35000]}

                    --- ASO (DOCUMENTO A SER AUDITADO) ---
                    {texto_aso[:15000]}

                    --- RESSALVAS E SOLICITAÇÃO DA AUDITORIA ---
                    {ressalvas_text_2}

                    ESTRUTURA OBRIGATÓRIA:
                    # PARECER TÉCNICO MÉDICO DE CONFORMIDADE ASO x PCMSO — NR-07
                    1. **Objeto e Fundamentação Médica:** Análise comparativa entre as diretrizes do PCMSO e os dados extraídos do ASO.
                    2. **Apontamento Específico de Divergências (ASO vs PCMSO):** Valide rigorosamente os exames que constam no ASO em relação ao PCMSO. Se os exames estão presentes e batem com a norma, ateste a conformidade. Aponte divergências reais apenas se houver ausência real ou periodicidade incorreta.
                    3. **Análise Crítica das Ressalvas:** Resposta técnica e fundamentada para a glosa/ressalva médica informada (Procedente / Improcedente).
                    4. **Conclusão e Providências:**
                       - **4.1. Veredito Geral (DESTAQUE MÁXIMO OBRIGATÓRIO):** Declaração explícita, em destaque, informando se o ASO está ou NÃO em conformidade com o PCMSO com base nos dados reais encontrados.

                    ## APÊNDICE – PLANO DE ADEQUAÇÃO E CRONOGRAMA DE EXAMES DO PCMSO
                    | Setor / Função | Exame / Procedimento | Diretriz NR-07 (PCMSO) | Status no ASO | Ação Corretiva Necessária | Prazo |
                    """

                    response = client.chat.completions.create(
                        model="deepseek-chat",
                        messages=[{"role": "system", "content": "Motor médico IA AuditGuard."}, {"role": "user", "content": prompt_2}],
                        stream=False,
                    )

                    resp_text = response.choices[0].message.content
                    bytes_word = converter_markdown_para_word_paisagem(resp_text)
                    incrementar_uso_auditoria(usuario_logado)

                    st.session_state.texto_resposta_2 = resp_text
                    st.session_state.docx_bytes_2 = bytes_word
                    st.success("✅ Auditoria Médica cruzada (PCMSO x ASO) gerada com sucesso!")
                except Exception as e:
                    st.error(f"❌ Erro: {str(e)}")

    if st.session_state.docx_bytes_2 and st.session_state.texto_resposta_2:
        st.markdown("---")
        st.download_button("📥 Baixar Parecer PCMSO x ASO (.docx)", data=st.session_state.docx_bytes_2, file_name="AuditGuard_PCMSO_x_ASO.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_2")
        st.write(st.session_state.texto_resposta_2)
